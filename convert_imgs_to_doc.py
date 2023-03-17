#!/usr/bin/env python3

import docx
import os

from docx.shared import Inches

rdir = '/Users/ahsannaveed/Desktop'
cpics = 'Compressed Pics'
apics = 'Annotated Pics'
compressed_files = os.listdir(os.path.join(rdir, cpics))

for i, f in enumerate(compressed_files):
    # ignore hidden files
    if not f.startswith('.'):
        img_path = os.path.join(rdir, cpics, f)
        doc = docx.Document()
        doc.add_picture(img_path, width=Inches(6), height=Inches(4))
        doc.add_paragraph(text='Name: FirstName LastName\nDOB: DOB\nContext: Event Name\nLeft-Right:')
        doc.save(os.path.join(rdir, apics, f'Ahsan-Nimra-Nikah-{i + 1}.docx'))

